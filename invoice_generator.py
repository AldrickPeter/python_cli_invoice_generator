import pandas as pd
import os
import docx
from docx import Document
import datetime

DEFAULT_PATH = ""

def create_file(df, filename):
    pd.DataFrame.to_excel(df, DEFAULT_PATH+filename)
    pass

def owner_create():
    print("Info file does not exist")
    proceed = input("To continue press 'y'\nif not press 'n'\n")
    if proceed == 'y':
        owner = pd.DataFrame(data={'info':['name', 'address', 'phone number'], 'values':[0, 0, 0]})
        name = input("Enter your business name: ")
        print("Enter your address: ")
        print("Enter 'over' once you finish the address")
        address = ""        
        while True:
            t = input()
            if t == "over":
                break
            else:
                address += t + '|'
        ph_number = input("Enter your phone number: ")
        owner["values"][0] = name
        owner["values"][1] = address
        owner["values"][2] = ph_number
        create_file(owner, "info.xlsx")


def add_new_customer():
    if os.path.isdir(DEFAULT_PATH + "customer") == False: 
        os.mkdir("customer")
    owner = pd.DataFrame(data={'info':['name', 'address', 'phone number'], 'values':[0, 0, 0]})
    name = input("Enter customer name: ")
    print("Enter customer address: ")
    print("Enter 'over' once you finish the address")
    address = ""        
    while True:
        t = input()
        if t == "over":
            break
        else:
            address += t + '|'
    ph_number = input("Enter customer phone number: ")
    owner["values"][0] = name
    owner["values"][1] = address
    owner["values"][2] = ph_number
    create_file(owner, "customer/ " + name + ".xlsx")
    pass

def add_new_item():
    if os.path.isdir(DEFAULT_PATH + "item") == False: os.mkdir("item")
    owner = pd.DataFrame(data={'info':['name', 'value'], 'values':[0, 0]})
    name = input("Enter item name: ")
    value = input("Enter item price: ")
    owner["values"][0] = name
    owner["values"][1] = float(value)
    create_file(owner, "item/ " + name + ".xlsx")
    pass

def modify_customer_info():
    pass

def modify_item_info():
    pass

if __name__ == "__main__":
    if not os.path.isfile("info.xlsx"): owner_create()            
    owner = pd.read_excel("info.xlsx")
    address_own = owner["values"][1].replace('|', "\n")    
    if not os.path.isdir("invoices"):
        os.mkdir("invoices")
    while True:
        total = 0
        instructions = """
        To create a new customer press 1
        To create a new item press 2
        To create a bill press 3
        Press 'q' to quit"""

        print(instructions)
        mode = input("Enter your response: ")
        if mode == 'q': break
        if mode == '1': add_new_customer()
        if mode == '2': add_new_item()
        if mode == '3':
            invoice_number = input("Enter the invoice number: ")
            customers = os.listdir('customer')
            items = os.listdir('item')
            i = 0
            item_sel = []
            quantity = []
            for customer in customers:
                print(customer[:-5], i+1)
            customer_sel = int(input("Enter the customer ID: "))-1
            i = 0
            while True:
                for item in items:
                    print(item[:-5], i+1)
                t = input("Enter the item ID: ")
                if t is not 'x': 
                    item_sel.append(int(t)-1)
                    quantity.append(int(input("Enter the quantity: ")))
                    print("Enter 'x' to exit")
                else: break
            
            tax = int(input("Enter the tax percentage: "))
            customer_df = pd.read_excel("customer/" + customers[customer_sel])

            invoice = Document()
            invoice.add_heading("Invoice", 0)

            invoice.add_heading(owner['values'][0], 1)
            invoice.add_paragraph(address_own + "Phone Number: " + owner["values"][2])
            p_1 = invoice.add_paragraph()
            p_1.add_run("Bill To,\n").bold = True
            p_1.add_run(customer_df["values"][0] + "\n")
            p_1.add_run(customer_df["values"][1].replace('|', "\n"))
            p_1.add_run("Phone Number: " + customer_df["values"][2])

            invoice.add_paragraph("Date: ")#,str(datetime.date.today()))

            table = invoice.add_table(rows=1, cols=3)
            table.rows[0].cells[0].text = "Item"
            table.rows[0].cells[1].text = "Quantity"
            table.rows[0].cells[2].text = "Price"

            for i in range(len(item_sel)):
                t_df = pd.read_excel("item/"+items[item_sel[i]])
                row_cell = table.add_row().cells
                row_cell[0].text = t_df["values"][0]
                row_cell[1].text = str(quantity[i])
                t = t_df["values"][1] * quantity[i]
                total += t
                row_cell[2].text = str(t)
            
            tax_calculated = total*(tax/100)
            total += tax_calculated

            row_tax = table.add_row().cells
            row_tax[1].text = "Tax " + str(tax) + "%"
            row_tax[2].text = str(tax_calculated)
            
            row_total = table.add_row().cells
            row_total[1].text = "Total "
            row_total[2].text = str(total)


            invoice.save("invoices/"+str(invoice_number)+".docx")
            pass

    pass